from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_word_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('SafetySense Sistem Geliştirme ve Otomasyon Planı', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Bu döküman, mevcut sistemdeki takip (tracking) hatalarından kaynaklanan mükerrer kayıtları engellemek, veritabanı temizliğini otomatize etmek ve bildirim sistemini kurmak için yapılacak değişiklikleri özetler.')
    
    # Section 1
    doc.add_heading('1. Mükerrer Kayıtların Önlenmesi (Cooldown Mekanizması)', level=1)
    doc.add_paragraph('Nesne takibi sırasında (YOLO Tracking) ID değişimi yaşansa bile, aynı ihlalin peş peşe defalarca kaydedilmesini engellemek için her kamera için bir "soğuma süresi" (cooldown) eklenecektir.')
    p1 = doc.add_paragraph()
    p1.add_run('• Mantık: ').bold = True
    p1.add_run('Bir ihlal loglandığında, o kamera için kayıt fonksiyonu 20 saniye boyunca kilitlenir.')
    p2 = doc.add_paragraph()
    p2.add_run('• Sonuç: ').bold = True
    p2.add_run('Aynı araç veya kişi için sadece bir kez mail gider ve DB kaydı oluşur.')
    
    # Section 2
    doc.add_heading('2. E-Posta Bildirim Sistemi Yapılandırması', level=1)
    doc.add_paragraph('İhlal anında ilgili kişiye (İSG Uzmanı/Müdür) anlık mail gitmesi için mailer.py dosyası kullanılacaktır.')
    p3 = doc.add_paragraph()
    p3.add_run('• Gereksinim: ').bold = True
    p3.add_run('Gmail üzerinden gönderim yapılacaksa "Uygulama Şifresi" (App Password) alınmalıdır.')
    
    # Section 3
    doc.add_heading('3. Otomatik Veri Temizliği (3 Günlük Döngü)', level=1)
    doc.add_paragraph('Sistemin şişmesini önlemek için her gece (veya sistem açılışında) 3 günden eski olan kayıtlar (DB satırları ve ilgili resim/video dosyaları) otomatik olarak silinecektir.')
    
    # Section 4
    doc.add_heading('4. Dashboard Optimizasyonu', level=1)
    doc.add_paragraph('Müdürün izleyeceği ekranda sadece en güncel ve temiz "screenshot" (kanıt fotoğrafları) listelenecek şekilde server.py endpointleri güncellenecektir.')
    
    doc.add_heading('İletişim ve Notlar', level=1)
    doc.add_paragraph('Bu değişiklikler mevcut analiz mantığını bozmadan sadece "kayıt ve bildirim" katmanını iyileştirecektir.')
    
    doc.save('Sistem_Gelistirme_Plani.docx')
    print("Sistem_Gelistirme_Plani.docx başarıyla oluşturuldu.")

if __name__ == "__main__":
    create_word_doc()
